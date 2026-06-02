"""Text extraction and metadata derivation (ADR-0007, module map §core/extract).

Single chokepoint for "PDF/DOCX/ODT → text + offsets" and "text → abstract/
keywords/fecha suggestions". Owns the per-format dispatch, the OCR gate
threshold, the abstract regex, the YAKE configuration, the fecha cover-page
heuristic, and the encrypted-PDF probe.
"""

from __future__ import annotations

import io
import json
import logging
import re
from dataclasses import dataclass, field
from datetime import date

import httpx

from buscasam.core import blob_store
from buscasam.settings import settings

logger = logging.getLogger(__name__)

_PDF_MIME = "application/pdf"
_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_ODT_MIME = "application/vnd.oasis.opendocument.text"

# ADR-0007 §4: PDF OCR gate threshold.
OCR_MIN_CHARS_PER_PAGE = 100


class PDFEncryptionError(Exception):
    pass


class OCRRequired(Exception):
    def __init__(self, sha256: str) -> None:
        self.sha256 = sha256


@dataclass(frozen=True)
class ExtractedDoc:
    text: str
    paragraph_breaks: list[int]
    page_breaks: list[int]
    raw_metadata: dict = field(default_factory=dict)


@dataclass(frozen=True)
class IndexableMetadata:
    abstract: str
    keywords: list[str]
    fecha: date | None


def probe_encrypted(data: bytes) -> None:
    """ADR-0007 §9: raise PDFEncryptionError if `data` is a password-protected PDF.

    Uses pdfminer to detect the encryption dictionary properly. Non-encryption
    parse failures (corrupted, truncated) are deferred to async indexing per
    ADR-0007 §9 and do not surface here.
    """
    from pdfminer.pdfdocument import PDFDocument
    from pdfminer.pdfdocument import PDFPasswordIncorrect as _PDFPwd
    from pdfminer.pdfparser import PDFParser

    try:
        PDFDocument(PDFParser(io.BytesIO(data)))
    except _PDFPwd as e:
        raise PDFEncryptionError("PDF is password-protected") from e
    except Exception:
        return


async def _read_blob(sha256: str) -> bytes:
    buf = bytearray()
    async for chunk in blob_store.open_for_send(sha256):
        buf.extend(chunk)
    return bytes(buf)


def _build_doc_from_paragraphs(paragraphs: list[str]) -> ExtractedDoc:
    pieces: list[str] = []
    breaks: list[int] = []
    cursor = 0
    for p in paragraphs:
        if not p.strip():
            continue
        pieces.append(p)
        cursor += len(p)
        breaks.append(cursor)
        pieces.append("\n\n")
        cursor += 2
    text = "".join(pieces).rstrip()
    # Drop the final break if it lands past the rstrip
    breaks = [b for b in breaks if b <= len(text)]
    return ExtractedDoc(
        text=text, paragraph_breaks=breaks, page_breaks=[], raw_metadata={}
    )


def _extract_docx(data: bytes) -> ExtractedDoc:
    from docx import Document as DocxDocument

    docx = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text for p in docx.paragraphs]
    return _build_doc_from_paragraphs(paragraphs)


def _extract_odt(data: bytes) -> ExtractedDoc:
    from odf.opendocument import load
    from odf.text import H, P
    from odf.teletype import extractText

    odt = load(io.BytesIO(data))
    paragraphs: list[str] = []
    for node in odt.getElementsByType(P) + odt.getElementsByType(H):
        paragraphs.append(extractText(node))
    return _build_doc_from_paragraphs(paragraphs)


def _extract_pdf(data: bytes) -> tuple[ExtractedDoc, int]:
    """Returns (doc, page_count). page_count includes blank pages (for OCR gate)."""
    from pdfminer.high_level import extract_text
    from pdfminer.pdfdocument import PDFDocument
    from pdfminer.pdfparser import PDFParser

    parser = PDFParser(io.BytesIO(data))
    pdfdoc = PDFDocument(parser)
    raw_metadata = (
        getattr(pdfdoc, "info", [{}])[0] if getattr(pdfdoc, "info", None) else {}
    )

    full_text = extract_text(io.BytesIO(data)) or ""
    # pdfminer separates pages with form feed \x0c, trailing one after the last page.
    pages = full_text.split("\x0c")
    if pages and pages[-1] == "":
        pages = pages[:-1]
    page_count = max(1, len(pages))

    pieces: list[str] = []
    paragraph_breaks: list[int] = []
    page_breaks: list[int] = []
    cursor = 0
    for page_idx, page in enumerate(pages):
        for para in page.split("\n\n"):
            stripped = para.strip()
            if not stripped:
                continue
            pieces.append(stripped)
            cursor += len(stripped)
            paragraph_breaks.append(cursor)
            pieces.append("\n\n")
            cursor += 2
        if page_idx < len(pages) - 1:
            page_breaks.append(cursor)

    text = "".join(pieces).rstrip()
    paragraph_breaks = [b for b in paragraph_breaks if b <= len(text)]
    page_breaks = [b for b in page_breaks if 0 < b <= len(text)]

    return (
        ExtractedDoc(
            text=text,
            paragraph_breaks=paragraph_breaks,
            page_breaks=page_breaks,
            raw_metadata=raw_metadata,
        ),
        page_count,
    )


_ABSTRACT_HEADING = re.compile(
    r"^(?:Resumen|Abstract|Summary|Sinopsis)\b",
    re.IGNORECASE | re.MULTILINE,
)
_NEXT_HEADING = re.compile(
    r"^(?:Introducci[oó]n|Cap[ií]tulo|Objetivos|Marco te[oó]rico|Metodolog[ií]a|"
    r"Conclusiones|Referencias|Bibliograf[ií]a|Index|Contenidos?)\b",
    re.IGNORECASE | re.MULTILINE,
)
# First real content heading — everything before it is cover/index/acknowledgements.
_FRONT_MATTER_END = re.compile(
    r"^(?:Resumen|Abstract|Summary|Sinopsis|Introducci[oó]n)\b",
    re.IGNORECASE | re.MULTILINE,
)
# Closing section that carries the findings (results/conclusions/discussion).
_CONCLUSION_HEADING = re.compile(
    r"^(?:Conclusi[oó]n(?:es)?|Resultados|Discusi[oó]n|Recomendaciones)\b",
    re.IGNORECASE | re.MULTILINE,
)
# Any section heading — used to bound a section's body at the next heading.
_ANY_HEADING = re.compile(
    r"^(?:Resumen|Abstract|Summary|Sinopsis|Introducci[oó]n|Cap[ií]tulo|"
    r"Objetivos|Marco te[oó]rico|Metodolog[ií]a|Materiales|M[eé]todos|"
    r"Resultados|Discusi[oó]n|Conclusi[oó]n(?:es)?|Recomendaciones|"
    r"Referencias|Bibliograf[ií]a|Anexos?|Ap[eé]ndices?|Index|Contenidos?)\b",
    re.IGNORECASE | re.MULTILINE,
)
_ABSTRACT_WORD_CAP = 300
_LLM_TEXT_CHAR_CAP = 12000
# Budget reserved for the closing section (conclusions/results) when the body is
# too long to send whole: the contribution lives at the end, which a head-only
# slice never reaches.
_LLM_CONCLUSION_CHAR_CAP = 5000
# Don't strip "front matter" if the first content heading appears only deep into
# the document — past this point a missing heading means odd formatting, not a
# real cover/index to drop.
_FRONT_MATTER_MAX_SKIP = 20000
_KEYWORD_CAP = 10
_KEYWORD_BLOCKLIST = {
    "este trabajo",
    "el presente trabajo",
    "presente trabajo",
    "presente informe",
    "este informe",
    "trabajo práctico",
    "trabajo practico",
    "la presente investigación",
    "presente investigación",
    "presente investigacion",
    "este documento",
    "el objetivo",
    "objetivo general",
    "universidad nacional",
    "universidad nacional de san martín",
    "universidad nacional de san martin",
    "san martín",
    "san martin",
}
_METADATA_LLM_SCHEMA = {
    "type": "object",
    "properties": {
        "abstract": {"type": "string"},
        "keywords": {
            "type": "array",
            "items": {"type": "string"},
            "minItems": 0,
            "maxItems": _KEYWORD_CAP,
        },
    },
    "required": ["abstract", "keywords"],
    "additionalProperties": False,
}
# Gemini's response_schema rejects `additionalProperties`; same shape otherwise.
_VERTEX_RESPONSE_SCHEMA = {
    "type": "object",
    "properties": {
        "abstract": {"type": "string"},
        "keywords": {"type": "array", "items": {"type": "string"}},
    },
    "required": ["abstract", "keywords"],
}
_PORTUGUESE_MARKERS = re.compile(
    r"\b(este documento descreve|previs[aã]o|s[eé]ries temporais|redes neurais|"
    r"avalia[cç][aã]o|j[aá]|por outro lado|informa[cç][oõ]es|aprendizagem)\b",
    re.IGNORECASE,
)

_COVER_TOKENS = re.compile(
    r"\b(tesis|tesina|trabajo|presentado|defendido|publicado)\b",
    re.IGNORECASE,
)
_YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")


def _truncate_words(s: str, cap: int) -> str:
    parts = s.split()
    if len(parts) <= cap:
        return " ".join(parts)
    return " ".join(parts[:cap])


def _normalize_phrase(s: str) -> str:
    cleaned = re.sub(r"\s+", " ", s).strip(" \t\r\n.,;:()[]{}\"'")
    return cleaned


def _clean_keywords(keywords: list[str]) -> list[str]:
    cleaned: list[str] = []
    seen: set[str] = set()
    for raw in keywords:
        phrase = _normalize_phrase(str(raw))
        key = phrase.casefold()
        if (
            not phrase
            or key in seen
            or key in _KEYWORD_BLOCKLIST
            or any(noise in key for noise in _KEYWORD_BLOCKLIST)
        ):
            continue
        if len(phrase.split()) > 5:
            continue
        seen.add(key)
        cleaned.append(phrase)
        if len(cleaned) >= _KEYWORD_CAP:
            break
    return cleaned


def _derive_abstract(text: str) -> str:
    if not text.strip():
        return ""
    head = text[:8000]  # first ~2 pages worth
    explicit = _derive_explicit_abstract(head)
    if explicit is not None:
        return explicit
    # fallback: first 1-3 paragraphs
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not paragraphs:
        return ""
    return _truncate_words(" ".join(paragraphs[:3]), _ABSTRACT_WORD_CAP)


def _derive_explicit_abstract(head: str) -> str | None:
    m = _ABSTRACT_HEADING.search(head)
    if not m:
        return None
    body = head[m.end() :]
    nxt = _NEXT_HEADING.search(body)
    body = body[: nxt.start()] if nxt else body
    return _truncate_words(body.strip(), _ABSTRACT_WORD_CAP)


def _derive_keywords(text: str) -> list[str]:
    if not text.strip():
        return []
    try:
        from yake import KeywordExtractor

        kw = KeywordExtractor(lan="es", n=3, dedupLim=0.7, top=8)
        results = kw.extract_keywords(text)
        return _clean_keywords([phrase for phrase, _score in results])
    except Exception:
        # YAKE failure must not block indexing (keywords are best-effort
        # suggestions per ADR-0007 §7); log so operators see degraded output.
        logger.warning("yake_failed", exc_info=True)
        return []


_PDF_CREATION_DATE_RE = re.compile(r"^D:(\d{4})")


def _derive_fecha_from_text(text: str) -> date | None:
    head = text[:8000]
    current_year = date.today().year
    best: int | None = None
    for m in _YEAR_RE.finditer(head):
        year = int(m.group())
        if not (1970 <= year <= current_year + 1):
            continue
        # check cover-token proximity (within 80 chars window before)
        window_start = max(0, m.start() - 80)
        if _COVER_TOKENS.search(head[window_start : m.end()]):
            if best is None or year > best:
                best = year
    return date(best, 1, 1) if best else None


def _derive_fecha_from_metadata(raw_metadata: dict) -> date | None:
    """ADR-0007 §8 step 2: fall back to PDF `/CreationDate` if plausible.

    PDF dates use `D:YYYYMMDDHHmmSS+TZ`; values may be `str` or `bytes`.
    """
    if not raw_metadata:
        return None
    raw = raw_metadata.get("CreationDate") or raw_metadata.get(b"CreationDate")
    if raw is None:
        return None
    if isinstance(raw, bytes):
        try:
            raw = raw.decode("ascii", errors="ignore")
        except Exception:
            return None
    m = _PDF_CREATION_DATE_RE.match(raw)
    if not m:
        return None
    year = int(m.group(1))
    current_year = date.today().year
    if not (1970 <= year <= current_year + 1):
        return None
    return date(year, 1, 1)


def _derive_fecha(doc: ExtractedDoc) -> date | None:
    from_text = _derive_fecha_from_text(doc.text)
    if from_text is not None:
        return from_text
    return _derive_fecha_from_metadata(doc.raw_metadata)


def derive_metadata(doc: ExtractedDoc) -> IndexableMetadata:
    return IndexableMetadata(
        abstract=_derive_abstract(doc.text),
        keywords=_derive_keywords(doc.text),
        fecha=_derive_fecha(doc),
    )


@dataclass(frozen=True)
class _LlmMetadata:
    abstract: str
    keywords: list[str]


_SPANISH_SYSTEM_INSTRUCTION = (
    "Respondé siempre en español, sin excepción. Nunca uses portugués ni inglés. "
    "Traducí al español cualquier término del texto fuente que esté en otro idioma."
)


def _strip_front_matter(text: str) -> str:
    m = _FRONT_MATTER_END.search(text)
    if m and m.start() <= _FRONT_MATTER_MAX_SKIP:
        return text[m.start() :]
    return text


def _section_after(text: str, start: int) -> str:
    """Body of the section whose heading starts at `start`, up to the next heading."""
    body = text[start:]
    nl = body.find("\n")
    rest = body[nl + 1 :] if nl != -1 else ""
    nxt = _ANY_HEADING.search(rest)
    return (rest[: nxt.start()] if nxt else rest).strip()


def _last_conclusion_match(text: str) -> re.Match[str] | None:
    m = None
    for m in _CONCLUSION_HEADING.finditer(text):
        pass
    return m


def _select_llm_text(text: str) -> str:
    """Pick the most summary-relevant slice within `_LLM_TEXT_CHAR_CAP`.

    Short docs go whole. For long ones we drop front matter and pair the
    intro-ward head with the closing section (results/conclusions), which a
    plain head slice never reaches.
    """
    if len(text) <= _LLM_TEXT_CHAR_CAP:
        return text
    body = _strip_front_matter(text)
    if len(body) <= _LLM_TEXT_CHAR_CAP:
        return body
    cm = _last_conclusion_match(body)
    if cm:
        conclusion = _section_after(body, cm.start())[:_LLM_CONCLUSION_CHAR_CAP]
        head_budget = _LLM_TEXT_CHAR_CAP - len(conclusion)
        if conclusion and cm.start() > head_budget:
            return f"{body[:head_budget].rstrip()}\n\n[…]\n\n{conclusion.strip()}"
    return body[:_LLM_TEXT_CHAR_CAP]


def _metadata_prompt(doc: ExtractedDoc, fallback: IndexableMetadata) -> str:
    return (
        "Sos un asistente que resume documentos académicos.\n"
        "Devolvé solo JSON válido con esta forma exacta: "
        '{"abstract": "string", "keywords": ["string"]}.\n'
        "El abstract es un resumen autocontenido del documento que permite "
        "entender de qué trata sin abrir el archivo:\n"
        "- 5 a 8 oraciones, entre 120 y 180 palabras.\n"
        "- Primera oración: tema y aporte central, entendible por sí sola.\n"
        "- Después: el problema que aborda, el método o enfoque, los datos o "
        "el caso de estudio, y el resultado o conclusión principal.\n"
        "- Prosa corrida, concreta y específica del documento; nada de relleno, "
        "portada, índice ni plantilla institucional.\n"
        "Keywords: 3 a 10 frases académicas específicas, sin nombres de "
        "plantilla institucional.\n"
        "Idioma obligatorio: español. No uses portugués ni inglés. Traduce "
        "términos del texto fuente al español cuando haga falta.\n"
        "No inventes datos que no estén en el texto.\n\n"
        f"Pistas (no las copies literalmente):\n"
        f"- Resumen heurístico: {fallback.abstract}\n"
        f"- Keywords candidatas: {', '.join(fallback.keywords)}\n\n"
        "Resumí a partir del texto entre delimitadores. Puede ser un extracto "
        "(inicio y conclusiones) de un documento más largo: resumí la obra "
        "completa, no solo el comienzo. No copies JSON, código ni tablas desde "
        "el texto fuente.\n"
        "<texto>\n"
        f"{_select_llm_text(doc.text)}\n"
        "</texto>"
    )


def _parse_llm_metadata(raw: str) -> _LlmMetadata:
    try:
        payload = json.loads(raw)
    except json.JSONDecodeError as e:
        raise ValueError("metadata LLM returned invalid JSON") from e
    if not isinstance(payload, dict):
        raise ValueError("metadata LLM returned non-object JSON")
    abstract = payload.get("abstract")
    keywords = payload.get("keywords")
    if not isinstance(abstract, str) or not isinstance(keywords, list):
        raise ValueError("metadata LLM returned invalid schema")
    if not all(isinstance(k, str) for k in keywords):
        raise ValueError("metadata LLM returned invalid keyword schema")
    return _LlmMetadata(
        abstract=_truncate_words(abstract.strip(), _ABSTRACT_WORD_CAP),
        keywords=_clean_keywords(keywords),
    )


def _looks_portuguese(value: str) -> bool:
    return bool(_PORTUGUESE_MARKERS.search(value))


async def _call_metadata_llm(
    client, doc: ExtractedDoc, fallback: IndexableMetadata
) -> _LlmMetadata:
    if settings.metadata_llm_provider == "vertex":
        return await _call_vertex(client, doc, fallback)
    return await _call_ollama(client, doc, fallback)


async def _call_ollama(
    client: httpx.AsyncClient, doc: ExtractedDoc, fallback: IndexableMetadata
) -> _LlmMetadata:
    response = await client.post(
        "/api/generate",
        json={
            "model": settings.metadata_llm_model,
            "prompt": _metadata_prompt(doc, fallback),
            "stream": False,
            "format": _METADATA_LLM_SCHEMA,
        },
        timeout=settings.metadata_llm_timeout_s,
    )
    response.raise_for_status()
    payload = response.json()
    raw = payload.get("response") if isinstance(payload, dict) else None
    if not isinstance(raw, str):
        raise ValueError("metadata LLM response missing response string")
    return _parse_llm_metadata(raw)


async def _call_vertex(
    client, doc: ExtractedDoc, fallback: IndexableMetadata
) -> _LlmMetadata:
    try:
        response = await client.aio.models.generate_content(
            model=settings.metadata_llm_model,
            contents=_metadata_prompt(doc, fallback),
            config={
                "response_mime_type": "application/json",
                "response_schema": _VERTEX_RESPONSE_SCHEMA,
                "temperature": 0,
                "system_instruction": _SPANISH_SYSTEM_INSTRUCTION,
            },
        )
    except Exception as e:  # SDK raises its own error types; any failure -> fallback.
        raise ValueError("vertex metadata call failed") from e
    raw = getattr(response, "text", None)
    if not isinstance(raw, str):
        raise ValueError("metadata LLM response missing response string")
    return _parse_llm_metadata(raw)


def _new_metadata_client():
    if settings.metadata_llm_provider == "vertex":
        from google import genai

        return genai.Client(
            vertexai=True,
            project=settings.vertex_project,
            location=settings.vertex_location,
        )
    return httpx.AsyncClient(base_url=settings.metadata_llm_url)


async def _close_metadata_client(client) -> None:
    if isinstance(client, httpx.AsyncClient):
        await client.aclose()


async def suggest_metadata(doc: ExtractedDoc, client=None) -> IndexableMetadata:
    """Best-effort staged metadata path.

    Heuristics always produce the fallback. The local LLM may clean up fallback
    output, but any timeout/outage/malformed output is non-fatal.
    """
    fallback = derive_metadata(doc)
    if not settings.metadata_llm_enabled or not doc.text.strip():
        return fallback

    owns_client = client is None
    if client is None:
        client = _new_metadata_client()
    try:
        llm = await _call_metadata_llm(client, doc, fallback)
    except (httpx.TimeoutException, httpx.HTTPError, ValueError):
        logger.warning("metadata_llm_failed", exc_info=True)
        return fallback
    finally:
        if owns_client:
            await _close_metadata_client(client)

    explicit = _derive_explicit_abstract(doc.text[:8000])
    abstract = explicit if explicit is not None else llm.abstract or fallback.abstract
    keywords = _clean_keywords(llm.keywords) or fallback.keywords
    if _looks_portuguese(" ".join([abstract, *keywords])):
        logger.warning("metadata_llm_non_spanish")
        return fallback
    return IndexableMetadata(abstract=abstract, keywords=keywords, fecha=fallback.fecha)


async def extract(sha256: str, mime: str) -> ExtractedDoc:
    data = await _read_blob(sha256)

    if mime == _DOCX_MIME:
        return _extract_docx(data)
    if mime == _ODT_MIME:
        return _extract_odt(data)
    if mime == _PDF_MIME:
        doc, page_count = _extract_pdf(data)
        # ADR-0007 §4: if average chars/page < threshold, OCR is required.
        if len(doc.text) / page_count < OCR_MIN_CHARS_PER_PAGE:
            raise OCRRequired(sha256)
        return doc
    raise ValueError(f"Unsupported mime: {mime}")
